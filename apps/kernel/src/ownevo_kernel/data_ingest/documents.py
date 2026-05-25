"""Parse PDF / DOCX uploads into text plus structured metadata.

pypdf and python-docx are imported lazily (the `data-ingest` extra) so the
kernel core stays dependency-light.

What gets surfaced beyond raw text:
  * DOCX has real structure — heading-styled paragraphs become `sections`
    (with level), `doc.tables` become `tables`, and the core properties give a
    `title`.
  * PDF has no reliable structure, so `sections` is a best-effort heuristic
    (short, heading-shaped lines) and `tables` is empty; `title` comes from the
    document info dictionary when present.
"""

from __future__ import annotations

import io

from .models import ParsedDocument, UploadKind


class DocumentParseError(ValueError):
    """The uploaded bytes could not be parsed as the declared kind."""


def parse_document(data: bytes, kind: UploadKind) -> ParsedDocument:
    """Parse document bytes into text + structured metadata."""
    if kind is UploadKind.PDF:
        return _parse_pdf(data)
    if kind is UploadKind.DOCX:
        return _parse_docx(data)
    raise DocumentParseError(f"not a document kind: {kind}")  # pragma: no cover


def _looks_like_heading(line: str) -> bool:
    s = line.strip()
    if not (0 < len(s) <= 80):
        return False
    if s.endswith((".", ",", ";", ":")):
        return False
    # ALL-CAPS or Title Case short lines read as headings in plain PDF text.
    letters = [c for c in s if c.isalpha()]
    if not letters:
        return False
    if s.isupper():
        return True
    words = s.split()
    return len(words) <= 8 and all(w[:1].isupper() for w in words if w[:1].isalpha())


def _parse_pdf(data: bytes) -> ParsedDocument:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - dependency guard
        raise DocumentParseError(
            "the `data-ingest` extra (pypdf) is required to parse PDFs"
        ) from exc
    try:
        reader = PdfReader(io.BytesIO(data))
        page_texts = [(page.extract_text() or "") for page in reader.pages]
    except Exception as exc:
        raise DocumentParseError(f"could not parse PDF upload: {exc}") from exc

    text = "\n\n".join(page_texts)
    title = None
    try:
        if reader.metadata and reader.metadata.title:
            title = str(reader.metadata.title)
    except Exception:
        title = None

    sections = [
        {"level": 1, "heading": ln.strip()}
        for ln in text.splitlines()
        if _looks_like_heading(ln)
    ]
    metadata = {
        "title": title,
        "page_count": len(page_texts),
        "section_count": len(sections),
        "table_count": 0,
        "tables_supported": False,
    }
    return ParsedDocument(metadata=metadata, text=text, sections=sections, tables=[])


def _parse_docx(data: bytes) -> ParsedDocument:
    try:
        import docx
    except ImportError as exc:  # pragma: no cover - dependency guard
        raise DocumentParseError(
            "the `data-ingest` extra (python-docx) is required to parse DOCX"
        ) from exc
    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:
        raise DocumentParseError(f"could not parse DOCX upload: {exc}") from exc

    paragraphs = [p.text for p in document.paragraphs]
    text = "\n".join(paragraphs)

    sections: list[dict[str, object]] = []
    for p in document.paragraphs:
        style = (p.style.name or "") if p.style else ""
        if style.startswith("Heading") and p.text.strip():
            # "Heading 2" -> level 2; "Heading"/"Title" -> level 1.
            tail = style.removeprefix("Heading").strip()
            level = int(tail) if tail.isdigit() else 1
            sections.append({"level": level, "heading": p.text.strip()})

    tables: list[list[list[str]]] = [
        [[cell.text for cell in row.cells] for row in table.rows]
        for table in document.tables
    ]

    title = None
    try:
        if document.core_properties.title:
            title = str(document.core_properties.title)
    except Exception:
        title = None

    metadata = {
        "title": title,
        "paragraph_count": len(paragraphs),
        "section_count": len(sections),
        "table_count": len(tables),
        "tables_supported": True,
    }
    return ParsedDocument(metadata=metadata, text=text, sections=sections, tables=tables)


__all__ = ["DocumentParseError", "parse_document"]
